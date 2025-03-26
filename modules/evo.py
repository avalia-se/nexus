import gradio as gr
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words
import warnings  # Somente se você realmente quer suprimir avisos


# Suprimir todos os avisos durante a execuÃ§Ã£o do script
warnings.filterwarnings("ignore")

# Calcular a data mais recente do CUB fora da funÃ§Ã£o
df_cub = pd.read_excel('dados/TABELAS_evo.xlsx', sheet_name='CUB')
ultimo_cub = df_cub.columns[-1]

#funÃ§Ã£o para escrever por extenso o valor atribuÃ­do
def numero_por_extenso(numero):
    reais = int(numero)

    extenso_reais = num2words(reais, lang='pt_BR')

    if reais == 1:
        extenso_reais += ' real'
    else:
        extenso_reais += ' reais'

    return extenso_reais

# FunÃ§Ã£o para calcular o valor do imÃ³vel
def calcular_valor_imovel(
                          area=100, data_refer_str=None, data_const_str=None, tipo_cub="R 1-N (Res. Unifamiliar)",
                          data_cub_str=None, percentual_cub=1, est_custo_dir="-", BDI=22.5, BDI_tipo ="Arbitrado", fator_local = 1,
                          just_fator_local = "-", tipologia="APARTAMENTOS",
                          estado="B - entre novo e regular", VR=0.0, deprec="Arbitrada",
                          valor_terreno=0,
                          est_ter="Grau III",
                          fc=1, fc_just="Arbitrado"):

    # Carregar dados dos arquivos Excel
    df_cub = pd.read_excel('dados/TABELAS_evo.xlsx', sheet_name='CUB')
    df_vida = pd.read_excel('dados/TABELAS_evo.xlsx', sheet_name='VUTIL')
    df_dep = pd.read_excel('dados/TABELAS_evo.xlsx', sheet_name='DEP')
    df_estado_cons = pd.read_excel('dados/TABELAS_evo.xlsx', sheet_name='CONS')

    # Converter datas de entrada
    if data_refer_str is None or data_refer_str == "":
        data_refer = datetime.now()
    else:
        data_refer = datetime.strptime(data_refer_str, "%m/%Y")
    data_const = datetime.strptime(data_const_str, "%m/%Y")
    if data_cub_str is None or data_cub_str == "":
        data_cub = datetime.now()
    else:
        data_cub = datetime.strptime(data_cub_str, "%m/%Y")

    # Filtrar por tipo_cub
    df_tipo_cub = df_cub[df_cub['CÃ“DIGO'] == tipo_cub]

    # Obter o valor do CUB na coluna correspondente Ã  data
    valor_cub_column = data_cub.strftime("%m/%Y")
    valor_cub = df_tipo_cub.at[df_tipo_cub.index[0], valor_cub_column]

    # Idade
    idade = data_refer.toordinal() - data_const.toordinal()
    if idade > 1:
        idade = idade // 365
    else:
        idade = 1

    # Filtrar por %de Vida
    vdu = df_vida.loc[(df_vida['FINAL'] == tipologia)]

    # % de vida Ãºtil
    perc_vdu = (idade / vdu['VIDAUTIL']) * 100
    perc_vdu = int(round(perc_vdu, 0))

    if perc_vdu >= 100:
        percentual_vdu = 100
    elif perc_vdu < 2:
        percentual_vdu = 2
    else:
        percentual_vdu = perc_vdu

    # Filtrar conforme o % de vida Ãºtil
    df_conserv = df_dep.loc[df_dep['%deVida'] == percentual_vdu]

     # Converter o valor residual de string para nÃºmero (float)
    try:
        VR = float(VR.strip()) if isinstance(VR, str) and VR.strip() != "" else 0.0
    except ValueError:
        VR = 0.0
    # Obter da depreciaÃ§Ã£o
    coef_HH = float(df_conserv[estado]/100)
    coef_HH = round(coef_HH, 3)

    # Valor do "Kd"
    kd = VR + (1 - coef_HH) * (1 - VR)
    kd = round(kd, 3)
        
    # Converter o valor do terreno de string para nÃºmero (float)
    #try:
        #valor_terreno = float(valor_terreno.strip()) if isinstance(valor_terreno, str) and valor_terreno.strip() != "" else 0.0
    #except ValueError:
        #valor_terreno = 0.0     
        
    # CÃ¡lculos
    Valor_sem_deprec = round(area * valor_cub * percentual_cub * (1 + BDI / 100) * fator_local, 2)
    Valor_com_deprec = Valor_sem_deprec * kd
    Valor_com_deprec = round(Valor_com_deprec, 2)
    valor_imovel = round((valor_terreno + Valor_com_deprec) * float(fc), 0)


    # Atributo da coluna "cons" pelo qual vocÃª deseja fazer a correspondÃªncia
    atributo_desejado = estado
    # Encontrar o valor da coluna "obs" com base no atributo da coluna "cons"
    valor_obs = df_estado_cons.loc[df_estado_cons['cons'] == atributo_desejado, 'obs'].iloc[0]

    # -------------------------------------- #

    # GRAU DE FUNDAMENTAÃ‡ÃƒO CUSTOS

    # item_1 - Graus de FundamentaÃ§Ã£o (Estimativa do custo direto)
    if est_custo_dir == "ElaboraÃ§Ã£o de orÃ§amento, no mÃ­nimo sintÃ©tico":
        cust_1 = 3
    elif est_custo_dir == "UtilizaÃ§Ã£o CUB para projeto semelhante ao projeto padrÃ£o":
        cust_1 = 2
    else:
        cust_1 = 1

    # item_2 - Graus de FundamentaÃ§Ã£o (BDI)
    if BDI_tipo == "Calculado":
        cust_2 = 3
    elif BDI_tipo == "Justificado":
        cust_2 = 2
    else:
        cust_2 = 1

    # item_3 - Graus de FundamentaÃ§Ã£o (DepreciaÃ§Ã£o fÃ­sica)
    if deprec == "Por levantamento do custo de recuperaÃ§Ã£o do bem, para deixÃ¡-lo no estado de novo ou caso de bens novos ou projetos hipotÃ©ticos":
        cust_3 = 3
    elif deprec == "Por mÃ©todos tÃ©cnicos consagrados, considerando-se idade, vida Ãºtil e estado de conservaÃ§Ã£o":
        cust_3 = 2
    else:
        cust_3 = 1

    # enquadramento
    soma_custo = cust_1 + cust_2 + cust_3
    if soma_custo >= 7 and cust_1 == 3 and cust_2 >= 2 and cust_3 >= 2:
        est_cr = "Grau III"
    elif soma_custo >= 5 and cust_1 >= 2 and cust_2 >= 2:
        est_cr = "Grau II"
    elif soma_custo >= 3 and cust_1 >= 1 and cust_2 >= 1 and cust_3 >= 1:
        est_cr = "Grau I"
    else:
        est_cr = "Fora dos critÃ©rios"

    # -------------------------------------- #

     # GRAU DE FUNDAMENTAÃ‡ÃƒO EVOLUTIVO

    # item_1 - Graus de FundamentaÃ§Ã£o (Estimativa do valor do terreno)
    if est_ter == "Grau III":
        evo_1 = 3
    elif est_ter == "Grau II":
        evo_1 = 2
    else:
        evo_1 = 1

    # item_2 - Graus de FundamentaÃ§Ã£o (Estimativa dos custos de reediÃ§Ã£o)
    if est_cr == "Grau III":
        evo_2 = 3
    elif est_cr == "Grau II":
        evo_2 = 2
    else:
        evo_2 = 1

    # item_3 - Graus de FundamentaÃ§Ã£o (Forma de cÃ¡lculo do FC)
    if fc_just == "Inferido em mercado semelhante":
        evo_3 = 3
    elif fc_just == "Justificado:":
        evo_3 = 2
    else:
        evo_3 = 1

    # enquadramento
    soma_evo = evo_1 + evo_2 + evo_3
    if soma_evo >= 8 and evo_1 == 3 and evo_2 == 3 and evo_3 >= 2: # confirmar!!!!!
        fundamentacao_evolutivo = "Grau III"
    elif soma_evo >= 5 and evo_1 >= 2 and evo_2 >= 2:
        fundamentacao_evolutivo = "Grau II"
    elif soma_evo >= 3 and evo_1 >= 1 and evo_2 >= 1 and evo_3 >= 1:
        fundamentacao_evolutivo = "Grau I"
    else:
        fundamentacao_evolutivo = "Fora dos critÃ©rios"

    # -------------------------------------- #

    # criaÃ§Ã£o de strings para os relatÃ³rios da interface e do word

    #####
    tipo_cub = tipo_cub.replace('.', '@').replace(',', '.').replace('@', ',')
    est_custo_dir = est_custo_dir.replace('.', '@').replace(',', '.').replace('@', ',')
    just_fator_local = just_fator_local.replace('.', '@').replace(',', '.').replace('@', ',')
    valor_inicial = f"""
    Ãrea construÃ­da : {area:,.2f} mÂ²
    Data de referÃªncia: {data_refer.strftime("%m/%Y")}
    Data da construÃ§Ã£o: {data_const_str}
    Data do CUB: {data_cub.strftime("%m/%Y")}
    Tipo de CUB: {tipo_cub}
    Fator para adequaÃ§Ã£o do CUB: {percentual_cub}
    Estimativa do custo direto: {est_custo_dir}
    BDI (%): {BDI}
    MÃ©todo de elaboraÃ§Ã£o do BDI: {BDI_tipo}
    Valor CUB: R$ {valor_cub:,.2f}
    Fator local: {fator_local}
    Fator local(justificativa): {just_fator_local}
    Valor antes da depreciaÃ§Ã£o = Ã¡rea construÃ­da * CUB * fator adequaÃ§Ã£o CUB * (1 + BDI / 100) * fator local
    Valor sem depreciaÃ§Ã£o: R$ {Valor_sem_deprec:,.2f}
    """

        # Substituindo ponto por vÃ­rgula
    valor_inicial = valor_inicial.replace('.', '@')
    valor_inicial = valor_inicial.replace(',', '.')
    valor_inicial = valor_inicial.replace('@', ',')

    #####
    valor_obs = valor_obs.replace('.', '@').replace(',', '.').replace('@', ',')
    deprec = deprec.replace('.', '@').replace(',', '.').replace('@', ',')
    deprec = f"""
    Tipologia: {tipologia}
    Vida Ãºtil da tipologia: {int(vdu['VIDAUTIL'])}
    Estado de conservaÃ§Ã£o: {estado}
    Estado de conservaÃ§Ã£o - descriÃ§Ã£o: {valor_obs}
    % de Vida Ãštil: {percentual_vdu}
    Coeficiente de DepreciaÃ§Ã£o: {coef_HH}
    Valor residual (0, 0.1 ou 0.2): {VR}
    Forma de cÃ¡lculo da depreciaÃ§Ã£o fÃ­sica: {deprec}
    Kd: {kd}
    onde: Kd = (Valor residual + (1 - percentual de depreciaÃ§Ã£o)*(1 - Valor residual))
    Valor depois da depreciaÃ§Ã£o = (Valor antes da depreciaÃ§Ã£o) x Kd (com coeficiente de valor residual)
    Valor final construÃ§Ã£o: R$ {Valor_com_deprec:,.2f}
    EspecificaÃ§Ã£o da AvaliaÃ§Ã£o (benfeitorias) - MÃ©todo da QuantificaÃ§Ã£o do Custo: {est_cr}
    """
    # Substituindo ponto por vÃ­rgula
    deprec = deprec.replace('.', '@')
    deprec = deprec.replace(',', '.')
    deprec = deprec.replace('@', ',')

    #####
    valor_ter = f"""
    Valor do Terreno: R$ {valor_terreno:,.2f}
    EspecificaÃ§Ã£o da AvaliaÃ§Ã£o (terreno ou gleba) -  MÃ©todo Comparativo ou Involutivo: {est_ter}
    """
    # Substituindo ponto por vÃ­rgula
    valor_ter = valor_ter.replace('.', '@')
    valor_ter = valor_ter.replace(',', '.')
    valor_ter = valor_ter.replace('@', ',')

    #####
    valor_ext = numero_por_extenso(valor_imovel)

    v_relat = f"{valor_imovel:,.2f}"
    v_relat = v_relat.replace('.', '@')
    v_relat = v_relat.replace(',', '.')
    v_relat = v_relat.replace('@', ',')

    valor_final = f"""
    FC (Fator de ComercializaÃ§Ã£o): {fc}
    ObservaÃ§Ã£o sobre o FC: {fc_just}
    ---------------------
    VI = (VT + CA) * FC
    Onde:
    VI: Valor estimado do imÃ³vel;
    VT: Valor estimado do terreno;
    CA: Custo de reediÃ§Ã£o das benfeitorias;
    FC: Fator de comercializaÃ§Ã£o.
    ---------------------
    Valor do ImÃ³vel (c/ arredondamento): R$ {v_relat}
    ({valor_ext})
    ---------------------
    EspecificaÃ§Ã£o da AvaliaÃ§Ã£o - MÃ©todo Evolutivo: {fundamentacao_evolutivo}
    """

    # -------------------------------------- #

    # CriaÃ§Ã£o de um relatÃ³rio da avaliaÃ§Ã£o no word
    # Criar um novo documento do Word
    doc = Document()

    # Definir o tÃ­tulo do documento
    doc.add_heading('RelatÃ³rio de AvaliaÃ§Ã£o de ImÃ³vel', level=1)

    # Definir as seÃ§Ãµes do relatÃ³rio
    sections = [
        (valor_inicial, "Valor Inicial da ConstruÃ§Ã£o"),
        (deprec, "CÃ¡lculo da DepreciaÃ§Ã£o"),
        (valor_ter, "Valor Estimado para o Terreno"),
        (valor_final, "Valor Final do ImÃ³vel"),
    ]

    for content, title in sections:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        run = p.add_run(content)
        run.font.size = Pt(12)

        if title == "":  # Se for a seÃ§Ã£o de assinatura
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Define o alinhamento para Ã  direita
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Define o alinhamento para Ã  esquerda

    # Salvar o documento em um arquivo .doc
    doc.save('relatorio_avaliacao.doc')

    # -------------------------------------- #

    # Outputs
    return valor_inicial, deprec, valor_ter, valor_final, 'relatorio_avaliacao.doc'

    # -------------------------------------- #

# Interface Gradio

def evo_tab():
    with gr.Tab("Evolutivo"):
        with gr.Row():
            with gr.Column():
                # Valor Inicial da ConstruÃ§Ã£o
                with gr.Row():
                    area = gr.Number(label="Ãrea (mÂ²)", value=100, info="Ãrea construÃ­da", scale=0.5)  # Remova a vÃ­rgula aqui
                    data_r = gr.Textbox(label="Data referÃªncia (mm/aaaa)", value=ultimo_cub, info="Data do fato gerador da demanda", scale=1)  # Remova a vÃ­rgula aqui
                    data_c = gr.Textbox(label="Data construÃ§Ã£o (mm/aaaa)", value=ultimo_cub, info="Data de construÃ§Ã£o do imÃ³vel", scale=1)
                with gr.Group():
                    with gr.Row():
                        cub = gr.Dropdown(label="Tipo de CUB", 
                                      choices=["R 1-B (Res. Unifamiliar)", "R 1-N (Res. Unifamiliar)", "R 1-A (Res. Unifamiliar)",
                                               "PP 4-B (PrÃ©dio Popular)", "PP 4-N (PrÃ©dio Popular)", "R 8-B  (Res. Multifamiliar)", 
                                               "R 8-N (Res. Multifamiliar)", "R 8-A (Res. Multifamiliar)", "R 16-N (Res. Multifamiliar)", 
                                               "R 16-A (Res. Multifamiliar)", "PIS (Projeto Inter. Social)", "RP1Q (ResidÃªncia Popular)", 
                                               "CAL 8-N (Com. Andar Livres)", "CAL 8-A (Com. Andar Livres)", 
                                               "CSL 8-N (Com.Salas e Lojas)", "CSL 8-A (Com.Salas e Lojas)", 
                                               "CSL 16-N (Com.Salas e Lojas)", "CSL 16-A (Com.Salas e Lojas)", "GI (GalpÃ£o Industrial)"], 
                                      value="R 1-N (Res. Unifamiliar)", scale=1)
                        data_cub = gr.Textbox(label="Data do CUB (mm/aaaa)", value=ultimo_cub, scale=1)
                        fator_cub = gr.Slider(0.5, 2.0, value=1, label="Fator para adequaÃ§Ã£o do CUB", 
                                          info="Para contemplar insumos/serviÃ§os que nÃ£o constam na composiÃ§Ã£o", step=0.1)
                    forma_cub = gr.Dropdown(["UtilizaÃ§Ã£o CUB para projeto semelhante ao projeto padrÃ£o",
                                          "UtilizaÃ§Ã£o CUB para projeto diferente do projeto padrÃ£o, com os devidos ajustes"], 
                                         label="Estimativa do custo direto", value="UtilizaÃ§Ã£o CUB para projeto semelhante ao projeto padrÃ£o")
                with gr.Row():
                    bdi = gr.Number(label="BDI (%)", value=22.5, scale=0.5)
                    tipo_bdi = gr.Dropdown(["Calculado", "Justificado", "Arbitrado"], label="Tipo de BDI", value="Justificado", scale=1)
                with gr.Row():
                    fator_local = gr.Slider(0.5, 1.5, value=1, label="Fator local", info="Atribua um coeficiente de valorizaÃ§Ã£o se necessÃ¡rio", step=0.1)
                    just_local = gr.Textbox(label="Justificativa para o Fator Local", value="-", info="Justifique tecnicamente o fator")
    
                # CÃ¡lculo da DepreciaÃ§Ã£o
                with gr.Group():
                    with gr.Row():
                        tipo = gr.Dropdown(label="Tipologia", choices=["APARTAMENTOS", "BANCOS", "CASAS DE ALVENARIA",
                                                                   "CASAS DE MADEIRA", "HOTÃ‰IS", "LOJAS", "TEATROS", 
                                                                   "ARMAZÃ‰NS", "FÃBRICAS", "CONST. RURAIS", "GARAGENS", 
                                                                   "EDIFÃCIOS DE ESCRITÃ“RIOS", "GALPÃ•ES (DEPÃ“SITOS)", "SILOS"], value="CASAS DE ALVENARIA", scale=1)
                        rst_con = gr.Dropdown(label="Estado de conservaÃ§Ã£o", choices=["A - novo", "B - entre novo e regular", 
                                                                                  "C - regular", "D - entre regular e reparos simples", 
                                                                                  "E - reparos simples", "F - entre reparos simples e importantes", 
                                                                                  "G - reparos importantes", "H - entre reparos importantes e sem valor"], value="A - novo", scale=1)
                        vr = gr.Dropdown(label="Valor residual", choices=["0", "0.1", "0.2"], value="0", scale=0.5)
                    deprec = gr.Dropdown(["Por levantamento do custo de recuperaÃ§Ã£o do bem, para deixÃ¡-lo no estado de novo ou caso de bens novos ou projetos hipotÃ©ticos", 
                                       "Por mÃ©todos tÃ©cnicos consagrados, considerando-se idade, vida Ãºtil e estado de conservaÃ§Ã£o", 
                                       "Arbitrado"], 
                                      label="DepreciaÃ§Ã£o FÃ­sica", 
                                      info="Forma pela qual a depreciaÃ§Ã£o fÃ­sica foi calculada", 
                                      value="Por mÃ©todos tÃ©cnicos consagrados, considerando-se idade, vida Ãºtil e estado de conservaÃ§Ã£o")
                #-------------------------------------------------------#
                button_1 = gr.Button("Calcular ConstruÃ§Ã£o")
                #-------------------------------------------------------#
                    
                # Valor do Terreno
                with gr.Row():
                    vt_inp = gr.Number(label="Valor do Terreno", info="Campo nÃ£o precisade de preenchimento p/ o cÃ¡lculo do Custo de ReproduÃ§Ã£o", value=0)
                    grau_t = gr.Radio(["Grau III", "Grau II", "Grau I"], info="Escolha o Grau de FundamentaÃ§Ã£o alcanÃ§ado pela avaliaÃ§Ã£o do terreno", label="", value="Grau III")
    
                # Valor final do imÃ³vel
                fc = gr.Slider(0.1, 2.0, value=1.0, label="FC (Fator de ComercializaÃ§Ã£o)", info="O Fator de ComercializaÃ§Ã£o, tambÃ©m chamado de Terceiro Componente ou Vantagem da Coisa Feita, Ã© definido no item 3.20 da NBR 14653-1:2001: â€œFator de comercializaÃ§Ã£o: RazÃ£o entre o valor de mercado de um bem e o seu custo de reediÃ§Ã£o ou de substituiÃ§Ã£o, que pode ser maior ou menor que 1", step=0.1)
                forma_fc = gr.Radio(["Inferido em mercado semelhante", "Justificado", "Arbitrado"], label="", info="Estipule como foi elaborado o FC", value="Justificado")
    
                #-------------------------------------------------------#
                button_2 = gr.Button("Calcular Terreno + ConstruÃ§Ã£o")
                #-------------------------------------------------------#
    
            with gr.Column():
                vi = gr.Textbox(label="Valor Inicial da ContruÃ§Ã£o")
                cd = gr.Textbox(label="CÃ¡lculo da DepreciaÃ§Ã£o")
                vt_out = gr.Textbox(label="Valor estimado para o terreno")
                vf = gr.Textbox(label="Valor final do imÃ³vel")
                la = gr.File(label="Laudo de AvaliaÃ§Ã£o")
    
        inputs = [
            area,              # Ãrea construÃ­da (mÂ²)
            data_r,            # Data de referÃªncia (mm/aaaa)
            data_c,            # Data da construÃ§Ã£o (mm/aaaa)
            cub,               # Tipo de CUB
            data_cub,          # Data do CUB (mm/aaaa)
            fator_cub,         # Fator para adequaÃ§Ã£o do CUB
            forma_cub,         # Estimativa do custo direto
            bdi,               # BDI (%)
            tipo_bdi,          # Tipo de BDI
            fator_local,       # Fator local
            just_local,        # Justificativa para o Fator Local
            tipo,              # Tipologia
            rst_con,           # Estado de conservaÃ§Ã£o
            vr,                # Valor residual (0, 0.1 ou 0.2)
            deprec,            # DepreciaÃ§Ã£o FÃ­sica
            vt_inp,                # Valor do Terreno
            grau_t,            # Grau de FundamentaÃ§Ã£o alcanÃ§ado pela avaliaÃ§Ã£o do terreno
            fc,                # FC (Fator de ComercializaÃ§Ã£o)
            forma_fc           # Forma como foi elaborado o FC
        ]
    
        outputs = [
            vi,  # Valor Inicial da ConstruÃ§Ã£o
            cd,  # CÃ¡lculo da DepreciaÃ§Ã£o
            vt_out,  # Valor estimado para o terreno
            vf,  # Valor final do imÃ³vel
            la   # Laudo de AvaliaÃ§Ã£o (arquivo)
        ]
    
        button_1.click(calcular_valor_imovel, inputs=inputs, outputs=outputs)
        button_2.click(calcular_valor_imovel, inputs=inputs, outputs=outputs)

    return locals()

# Substituir a tabela pelo cálculo da depreciação
